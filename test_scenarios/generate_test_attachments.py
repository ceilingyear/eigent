from __future__ import annotations

import csv
import json
import zipfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "attachments"


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def write_text(path: Path, content: str) -> None:
    path.write_text(content.strip() + "\n", encoding="utf-8")


def create_task1() -> None:
    task_dir = OUT / "task1_multiformat_research"
    ensure_dir(task_dir)

    # PDF source
    pdf_path = task_dir / "market_brief_2026.pdf"
    doc = SimpleDocTemplate(str(pdf_path), pagesize=A4)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("2026 AI Operations Market Brief", styles["Title"]),
        Spacer(1, 12),
        Paragraph(
            "This brief summarizes a fictional internal analysis of AI operations tools. "
            "The main finding is that teams gain the most value when agents are connected "
            "to controlled tools, review checkpoints, and durable execution logs.",
            styles["BodyText"],
        ),
        Spacer(1, 12),
    ]
    table_data = [
        ["Metric", "2025", "2026 Forecast", "Note"],
        ["Agent task completion", "58%", "71%", "Improved by better tool use"],
        ["Human review rate", "34%", "29%", "Lower but still required"],
        ["Average task runtime", "42 min", "55 min", "Longer workflows"],
    ]
    table = Table(table_data, colWidths=[130, 80, 100, 180])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.extend([table, Spacer(1, 12)])
    story.append(
        Paragraph(
            "Risk note: the spreadsheet from Finance uses a higher cost estimate than "
            "this brief. The final report should flag that conflict.",
            styles["BodyText"],
        )
    )
    doc.build(story)

    # DOCX source
    docx_path = task_dir / "meeting_notes.docx"
    word = Document()
    word.add_heading("AI Workflow Review Meeting Notes", level=1)
    word.add_paragraph("Date: 2026-06-10")
    word.add_paragraph("Attendees: Product, Engineering, Operations")
    word.add_heading("Decisions", level=2)
    for item in [
        "Prioritize file handling tests with PDF, DOCX, CSV, and screenshots.",
        "Require source links or file references in generated reports.",
        "Add a human confirmation step before sending external messages.",
    ]:
        word.add_paragraph(item, style="List Bullet")
    word.add_heading("Open Questions", level=2)
    word.add_paragraph(
        "Whether long-running workflows should be checkpointed after every tool call "
        "or only after each subtask completes."
    )
    word.save(docx_path)

    # CSV source
    csv_path = task_dir / "finance_assumptions.csv"
    with csv_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["category", "annual_cost_usd", "confidence", "owner"])
        writer.writerow(["model_api", 84000, "medium", "Finance"])
        writer.writerow(["browser_automation", 18000, "high", "Engineering"])
        writer.writerow(["document_processing", 26000, "medium", "Operations"])
        writer.writerow(["human_review", 42000, "low", "Operations"])

    # Image source
    img_path = task_dir / "workflow_snapshot.png"
    image = Image.new("RGB", (1000, 620), "#F7FAFC")
    draw = ImageDraw.Draw(image)
    try:
        font_title = ImageFont.truetype("arial.ttf", 34)
        font = ImageFont.truetype("arial.ttf", 22)
    except OSError:
        font_title = ImageFont.load_default()
        font = ImageFont.load_default()
    draw.text((40, 35), "Observed AI Workflow", fill="#111827", font=font_title)
    boxes = [
        ("User request", 60, 140),
        ("Planner", 285, 140),
        ("Tool calls", 510, 140),
        ("Validation", 735, 140),
        ("Report output", 510, 360),
    ]
    for label, x, y in boxes:
        draw.rounded_rectangle((x, y, x + 180, y + 90), radius=12, fill="#E0F2FE", outline="#0369A1", width=3)
        draw.text((x + 18, y + 32), label, fill="#0F172A", font=font)
    for start, end in [((240, 185), (285, 185)), ((465, 185), (510, 185)), ((690, 185), (735, 185)), ((600, 230), (600, 360))]:
        draw.line((*start, *end), fill="#334155", width=4)
    draw.text((60, 520), "Note: validation is missing in several failed runs.", fill="#B91C1C", font=font)
    image.save(img_path)

    write_text(
        task_dir / "prompt.md",
        """
        请读取我上传的全部附件，生成一份中文研究报告。要求：
        1. 分别总结 PDF、DOCX、CSV、PNG 中的关键信息。
        2. 找出文件之间的冲突点，尤其是成本估算和流程风险。
        3. 根据 CSV 生成成本汇总表。
        4. 输出 Markdown 报告，并另存一份 DOCX 或 PDF。
        5. 报告必须包含：摘要、资料来源、关键发现、冲突与风险、建议、附录。
        """,
    )


def create_task2() -> None:
    task_dir = OUT / "task2_dirty_excel_cleanup"
    ensure_dir(task_dir)
    xlsx_path = task_dir / "sales_ops_dirty_data.xlsx"

    wb = Workbook()
    ws = wb.active
    ws.title = "Orders"
    rows = [
        ["order_id", "date", "region", "sales_rep", "customer", "amount", "status"],
        ["A-1001", "2026/01/05", "East", "Lina", "Orion Labs", "$1,240.00", "Paid"],
        ["A-1002", "01-07-2026", "West", "Marco", "BluePeak", "980", "Paid"],
        ["A-1003", "2026.01.09", "North", "Ava", "Northstar", "", "Pending"],
        ["A-1004", "2026/01/12", "East", "Lina", "Orion Labs", "$1,240.00", "Paid"],
        ["A-1004", "2026/01/12", "East", "Lina", "Orion Labs", "$1,240.00", "Paid"],
        ["A-1005", "13/01/2026", "South", "Nate", "GreenWorks", "-300", "Refund"],
        ["A-1006", "2026/02/01", "west", "Marco", "BluePeak", "1,870 USD", "Paid"],
        ["A-1007", "invalid-date", "North", "", "Nova Retail", "2200", "Paid"],
        ["A-1008", "2026/02/14", "South", "Nate", "GreenWorks", "N/A", "Pending"],
    ]
    for row in rows:
        ws.append(row)
    ws.freeze_panes = "A2"
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E79")
    widths = [14, 16, 12, 14, 18, 16, 14]
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[chr(64 + idx)].width = width

    ws2 = wb.create_sheet("Targets")
    ws2.append(["region", "monthly_target", "owner"])
    ws2.append(["East", 5000, "Ops"])
    ws2.append(["West", 4800, "Ops"])
    ws2.append(["North", 4400, "Ops"])
    ws2.append(["South", 3900, "Ops"])
    for cell in ws2[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="375623")

    ws3 = wb.create_sheet("Notes")
    ws3.append(["Known issues"])
    ws3.append(["Duplicate order A-1004 appears twice."])
    ws3.append(["Date formats are inconsistent."])
    ws3.append(["Some amounts include currency symbols or text."])
    ws3.append(["Some region labels differ by case."])
    wb.save(xlsx_path)

    write_text(
        task_dir / "prompt.md",
        """
        请清洗并分析这个 Excel 文件。要求：
        1. 识别重复订单、缺失值、异常金额、非法日期和区域命名不一致问题。
        2. 输出一份数据质量问题清单。
        3. 生成清洗后的 Excel 文件，保留原始数据 sheet，并新增 Cleaned、Issue Log、Summary 三个 sheet。
        4. 统计各区域销售额、订单数、退款金额、完成率。
        5. 生成 Markdown 分析报告，说明清洗规则和关键发现。
        """,
    )


def create_task3() -> None:
    task_dir = OUT / "task3_code_logs_long_task"
    ensure_dir(task_dir)
    project_dir = task_dir / "sample_service"
    ensure_dir(project_dir / "src")
    ensure_dir(project_dir / "tests")

    write_text(
        project_dir / "README.md",
        """
        # Sample Service

        This is a small intentionally flawed service used to test long-task code analysis.
        It exposes a user import function and a basic retry helper.

        Expected behavior:
        - Import users from JSON lines.
        - Skip invalid rows without crashing.
        - Retry transient network errors up to 3 times.
        - Write a summary report.
        """,
    )
    write_text(
        project_dir / "src" / "service.py",
        """
        import json
        import time


        def parse_user_line(line):
            data = json.loads(line)
            return {
                "id": data["id"],
                "email": data["email"].lower(),
                "plan": data.get("plan", "free"),
            }


        def import_users(lines):
            users = []
            errors = []
            for index, line in enumerate(lines):
                user = parse_user_line(line)
                users.append(user)
            return {"users": users, "errors": errors}


        def retry_request(fn, attempts=3):
            last_error = None
            for attempt in range(attempts):
                try:
                    return fn()
                except Exception as exc:
                    last_error = exc
                    time.sleep(attempt)
            raise last_error
        """,
    )
    write_text(
        project_dir / "tests" / "test_service.py",
        """
        import pytest
        from src.service import import_users, retry_request


        def test_import_users_skips_invalid_rows():
            lines = [
                '{"id": "u1", "email": "A@EXAMPLE.COM", "plan": "pro"}',
                '{"id": "u2", "plan": "free"}',
                'not-json',
            ]
            result = import_users(lines)
            assert len(result["users"]) == 1
            assert len(result["errors"]) == 2
            assert result["users"][0]["email"] == "a@example.com"


        def test_retry_request_eventually_succeeds():
            calls = {"count": 0}

            def flaky():
                calls["count"] += 1
                if calls["count"] < 3:
                    raise RuntimeError("temporary")
                return "ok"

            assert retry_request(flaky, attempts=3) == "ok"
        """,
    )
    write_text(
        project_dir / "requirements.txt",
        """
        pytest==8.2.0
        """,
    )
    write_text(
        task_dir / "runtime.log",
        """
        2026-06-18T09:12:03Z INFO import started file=users.jsonl
        2026-06-18T09:12:04Z ERROR row=18 error=KeyError('email')
        2026-06-18T09:12:04Z ERROR row=19 error=JSONDecodeError('Expecting value')
        2026-06-18T09:12:06Z WARNING retry attempt=1 service=crm reason=temporary
        2026-06-18T09:12:07Z WARNING retry attempt=2 service=crm reason=temporary
        2026-06-18T09:12:09Z INFO import completed imported=148 failed=2 duration_ms=6210
        """,
    )
    write_text(
        task_dir / "requirements_change_request.md",
        """
        # Change Request

        The importer currently fails the whole job when one row is malformed.
        Please update it so invalid rows are collected into `errors` and valid rows continue.

        Acceptance criteria:
        - Missing required fields should not crash the import.
        - Invalid JSON should not crash the import.
        - Email should be normalized to lowercase.
        - Tests should pass.
        - Add a short summary of the root cause and fix.
        """,
    )

    zip_path = task_dir / "sample_service.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in project_dir.rglob("*"):
            zf.write(path, path.relative_to(task_dir))

    write_text(
        task_dir / "prompt.md",
        """
        请分析并修复这个压缩包里的 Python 小项目。要求：
        1. 解压 sample_service.zip，阅读 README、源码、测试和 runtime.log。
        2. 找出导致导入任务失败的根因。
        3. 修改代码，使无效行被记录到 errors，不能让整个任务崩溃。
        4. 运行测试并修复失败。
        5. 输出一份 Markdown 修复报告，包括根因、修改内容、测试结果、剩余风险。
        6. 这个任务请完整执行，不要只给建议。
        """,
    )


def create_index() -> None:
    ensure_dir(OUT)
    manifest = {
        "task1_multiformat_research": [
            "market_brief_2026.pdf",
            "meeting_notes.docx",
            "finance_assumptions.csv",
            "workflow_snapshot.png",
            "prompt.md",
        ],
        "task2_dirty_excel_cleanup": [
            "sales_ops_dirty_data.xlsx",
            "prompt.md",
        ],
        "task3_code_logs_long_task": [
            "sample_service.zip",
            "runtime.log",
            "requirements_change_request.md",
            "prompt.md",
        ],
    }
    (OUT / "manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    write_text(
        OUT / "README.md",
        """
        # Eigent 文件处理与长任务测试附件

        ## Task 1 - 多格式资料汇总

        目的：测试 PDF、DOCX、CSV、PNG 混合附件读取、跨文件汇总、冲突识别和报告生成能力。

        附件目录：`task1_multiformat_research/`

        上传文件：
        - `market_brief_2026.pdf`
        - `meeting_notes.docx`
        - `finance_assumptions.csv`
        - `workflow_snapshot.png`

        Prompt 使用：`task1_multiformat_research/prompt.md`

        ## Task 2 - 脏 Excel 数据清洗

        目的：测试 Excel 多 sheet 读取、数据清洗、异常识别、生成新 Excel 和分析报告能力。

        附件目录：`task2_dirty_excel_cleanup/`

        上传文件：
        - `sales_ops_dirty_data.xlsx`

        Prompt 使用：`task2_dirty_excel_cleanup/prompt.md`

        ## Task 3 - 代码压缩包 + 日志长任务

        目的：测试 ZIP 解压、代码阅读、日志分析、代码修改、运行测试和修复报告生成能力。

        附件目录：`task3_code_logs_long_task/`

        上传文件：
        - `sample_service.zip`
        - `runtime.log`
        - `requirements_change_request.md`

        Prompt 使用：`task3_code_logs_long_task/prompt.md`
        """,
    )


def main() -> None:
    create_task1()
    create_task2()
    create_task3()
    create_index()
    print(f"Generated attachments under: {OUT}")


if __name__ == "__main__":
    main()
